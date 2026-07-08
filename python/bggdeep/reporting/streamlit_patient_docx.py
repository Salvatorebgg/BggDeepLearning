# -*- coding: utf-8 -*-

"""
BggDeepLearning Streamlit patient DOCX report module

File:
python/bggdeep/reporting/streamlit_patient_docx.py

Purpose:
1. Build a Word DOCX report for one current patient prediction
2. Return DOCX bytes for Streamlit download_button
3. Keep Streamlit page code cleaner

Clinical note:
This report is generated from simulated-data models and is for engineering
demonstration only. It must not be used for real clinical decisions.
"""

from __future__ import annotations

from io import BytesIO
from typing import Dict

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


def set_document_default_style(document: Document) -> None:
    """
    Set default Chinese-compatible Word style.

    We use Microsoft YaHei as the East Asian font because it usually exists
    on Windows and displays Chinese text clearly.
    """
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if style_name in styles:
            style = styles[style_name]
            style.font.name = "Microsoft YaHei"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_title(document: Document, title: str, subtitle: str) -> None:
    """
    Add centered title and subtitle.
    """
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.name = "Microsoft YaHei"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    subtitle_paragraph = document.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle_run = subtitle_paragraph.add_run(subtitle)
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.name = "Microsoft YaHei"
    subtitle_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    document.add_paragraph("")


def add_key_value_table(
    document: Document,
    title: str,
    data: Dict[str, object],
) -> None:
    """
    Add key-value table to Word.
    """
    document.add_heading(title, level=2)

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    header_cells[0].text = "项目"
    header_cells[1].text = "内容"

    for key, value in data.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(key)
        row_cells[1].text = "" if pd.isna(value) else str(value)

    document.add_paragraph("")


def add_dataframe_table(
    document: Document,
    title: str,
    df: pd.DataFrame,
    max_rows: int = 20,
) -> None:
    """
    Add DataFrame as Word table.
    """
    document.add_heading(title, level=2)

    if df.empty:
        document.add_paragraph("当前没有可用数据。")
        return

    display_df = df.head(max_rows).copy()

    table = document.add_table(rows=1, cols=len(display_df.columns))
    table.style = "Table Grid"

    header_cells = table.rows[0].cells

    for index, column in enumerate(display_df.columns):
        header_cells[index].text = str(column)

    for _, row in display_df.iterrows():
        row_cells = table.add_row().cells

        for index, value in enumerate(row):
            if pd.isna(value):
                row_cells[index].text = ""
            elif isinstance(value, float):
                row_cells[index].text = f"{value:.6f}"
            else:
                row_cells[index].text = str(value)

    document.add_paragraph("")


def add_warning_section(document: Document) -> None:
    """
    Add clinical warning section.
    """
    document.add_heading("五、重要说明", level=2)

    notes = [
        "当前模型基于模拟数据训练，仅用于工程演示。",
        "当前预测结果不能作为真实临床诊疗依据。",
        "风险分层由页面中设置的低风险阈值和高风险阈值决定。",
        "SHAP 或线性贡献值用于解释模型行为，不代表因果关系。",
        "真实临床研究中需要进行外部验证、前瞻性验证和临床专家审核。",
    ]

    for item in notes:
        document.add_paragraph(item, style="List Bullet")


def build_patient_prediction_docx_bytes(
    model_name: str,
    probability: float,
    risk_group: str,
    predicted_label: int,
    raw_input_display: Dict[str, object],
    explanation_display_df: pd.DataFrame,
) -> bytes:
    """
    Build DOCX report and return bytes.

    Parameters
    ----------
    model_name:
        Chinese model display name.

    probability:
        Predicted probability of poor_outcome.

    risk_group:
        Chinese risk group.

    predicted_label:
        Binary predicted label.

    raw_input_display:
        Dictionary with Chinese feature labels and patient input values.

    explanation_display_df:
        Display-ready explanation DataFrame with Chinese columns.

    Returns
    -------
    bytes
        DOCX file bytes for Streamlit download_button.
    """
    document = Document()
    set_document_default_style(document)

    add_title(
        document=document,
        title="BggDeepLearning 当前患者风险预测报告",
        subtitle="Current Patient Risk Prediction Report",
    )

    model_info = {
        "模型名称": model_name,
        "预测结局": "poor_outcome",
        "报告类型": "单个患者预测报告",
        "数据来源说明": "模拟数据训练模型，仅用于工程演示",
    }

    prediction_info = {
        "预测风险概率": f"{probability:.6f}",
        "预测风险百分比": f"{probability:.1%}",
        "风险分层": risk_group,
        "预测标签": predicted_label,
    }

    add_key_value_table(
        document=document,
        title="一、模型信息",
        data=model_info,
    )

    add_key_value_table(
        document=document,
        title="二、预测结果",
        data=prediction_info,
    )

    add_key_value_table(
        document=document,
        title="三、患者输入信息",
        data=raw_input_display,
    )

    add_dataframe_table(
        document=document,
        title="四、主要解释性因素",
        df=explanation_display_df,
        max_rows=15,
    )

    add_warning_section(document)

    output_buffer = BytesIO()
    document.save(output_buffer)
    output_buffer.seek(0)

    return output_buffer.getvalue()