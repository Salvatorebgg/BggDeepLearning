"""
BggDeepLearning Word report generation module

File:
python/bggdeep/reporting/word_report.py

Purpose:
1. Generate a comprehensive DOCX report
2. Generate a manuscript-style clinical results DOCX report
3. Insert tables, summaries, and key figures
"""

from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Dict, List

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt


@dataclass
class WordReportConfig:
    """
    Word report configuration.
    """

    project_name: str = "BggDeepLearning"
    target_name: str = "poor_outcome"
    report_title: str = "Clinical Model Evaluation Report"
    author: str = "BggDeepLearning Team"


class WordModelReportBuilder:
    """
    Build Word reports from model evaluation outputs.
    """

    def __init__(self, config: WordReportConfig) -> None:
        self.config = config

    @staticmethod
    def read_csv_if_exists(path: Path) -> pd.DataFrame:
        if path.exists():
            return pd.read_csv(path)
        return pd.DataFrame()

    @staticmethod
    def set_default_styles(document: Document) -> None:
        """
        Set simple document styles.
        """
        styles = document.styles

        normal = styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(10.5)

        for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
            if style_name in styles:
                styles[style_name].font.name = "Times New Roman"

    @staticmethod
    def add_title_page(
        document: Document,
        title: str,
        subtitle: str,
        author: str,
    ) -> None:
        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(20)

        subtitle_paragraph = document.add_paragraph()
        subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_paragraph.add_run(subtitle)
        subtitle_run.font.size = Pt(13)

        document.add_paragraph("")

        author_paragraph = document.add_paragraph()
        author_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_paragraph.add_run(author)
        author_run.font.size = Pt(11)

        time_paragraph = document.add_paragraph()
        time_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        time_run = time_paragraph.add_run(
            f"Generated time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        )
        time_run.font.size = Pt(10)

        document.add_page_break()

    @staticmethod
    def add_dataframe_table(
        document: Document,
        df: pd.DataFrame,
        columns: List[str] | None = None,
        max_rows: int = 20,
        title: str | None = None,
    ) -> None:
        """
        Add pandas DataFrame as Word table.
        """
        if title:
            document.add_paragraph(title, style="Heading 3")

        if df.empty:
            document.add_paragraph("No available data.")
            return

        display_df = df.copy()

        if columns is not None:
            existing_columns = [column for column in columns if column in display_df.columns]
            if existing_columns:
                display_df = display_df[existing_columns]

        display_df = display_df.head(max_rows)

        table = document.add_table(rows=1, cols=len(display_df.columns))
        table.style = "Table Grid"

        header_cells = table.rows[0].cells
        for i, column in enumerate(display_df.columns):
            header_cells[i].text = str(column)

        for _, row in display_df.iterrows():
            cells = table.add_row().cells
            for i, value in enumerate(row):
                if pd.isna(value):
                    text = ""
                else:
                    text = str(value)
                cells[i].text = text

        document.add_paragraph("")

    @staticmethod
    def add_picture_if_exists(
        document: Document,
        image_path: Path,
        title: str,
        width_inches: float = 6.2,
    ) -> None:
        """
        Add image if it exists.
        """
        if not image_path.exists():
            document.add_paragraph(f"Figure not found: {image_path}")
            return

        document.add_paragraph(title, style="Heading 3")
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(width_inches))
        document.add_paragraph("")

    @staticmethod
    def add_bullet_list(document: Document, items: List[str]) -> None:
        for item in items:
            document.add_paragraph(item, style="List Bullet")

    def load_inputs(self, table_dir: Path) -> Dict[str, pd.DataFrame]:
        """
        Load all report input tables.
        """
        return {
            "best_model": self.read_csv_if_exists(table_dir / "model_best_model_summary.csv"),
            "leaderboard_validation": self.read_csv_if_exists(table_dir / "model_leaderboard_validation.csv"),
            "leaderboard_test": self.read_csv_if_exists(table_dir / "model_leaderboard_test.csv"),
            "roc_summary": self.read_csv_if_exists(table_dir / "model_comparison_roc_summary.csv"),
            "calibration_metrics": self.read_csv_if_exists(table_dir / "model_calibration_metrics.csv"),
            "dca_summary": self.read_csv_if_exists(table_dir / "dca_selected_threshold_summary.csv"),
            "file_index": self.read_csv_if_exists(table_dir / "comprehensive_model_report_file_index.csv"),
        }

    def build_comprehensive_docx(
        self,
        project_root: Path,
        table_dir: Path,
        figure_dir: Path,
        output_path: Path,
    ) -> Path:
        """
        Build comprehensive Word report.
        """
        inputs = self.load_inputs(table_dir)

        document = Document()
        self.set_default_styles(document)

        self.add_title_page(
            document=document,
            title="BggDeepLearning Comprehensive Model Evaluation Report",
            subtitle=f"Target outcome: {self.config.target_name}",
            author=self.config.author,
        )

        document.add_heading("1. Executive Summary", level=1)
        document.add_paragraph(
            "This report summarizes the current model development results, "
            "including discrimination, calibration, clinical net benefit, and model ranking. "
            "All results are generated from simulated clinical data and are intended for "
            "engineering demonstration rather than validated clinical decision-making."
        )

        document.add_heading("2. Current Best Model", level=1)
        self.add_dataframe_table(
            document,
            inputs["best_model"],
            columns=[
                "model_display_name",
                "model_name",
                "split",
                "auc",
                "accuracy",
                "sensitivity",
                "specificity",
                "precision",
                "recall",
                "f1",
                "balanced_summary_score",
            ],
            max_rows=5,
        )

        document.add_heading("3. Validation Leaderboard", level=1)
        self.add_dataframe_table(
            document,
            inputs["leaderboard_validation"],
            columns=[
                "rank_within_split",
                "model_display_name",
                "auc",
                "accuracy",
                "sensitivity",
                "specificity",
                "precision",
                "recall",
                "f1",
                "balanced_summary_score",
            ],
            max_rows=20,
        )

        document.add_heading("4. Test Leaderboard", level=1)
        self.add_dataframe_table(
            document,
            inputs["leaderboard_test"],
            columns=[
                "rank_within_split",
                "model_display_name",
                "auc",
                "accuracy",
                "sensitivity",
                "specificity",
                "precision",
                "recall",
                "f1",
                "balanced_summary_score",
            ],
            max_rows=20,
        )

        document.add_heading("5. ROC Curve Comparison", level=1)
        self.add_dataframe_table(
            document,
            inputs["roc_summary"],
            columns=["split", "model_name", "roc_auc"],
            max_rows=20,
            title="ROC summary",
        )

        self.add_picture_if_exists(
            document,
            figure_dir / "model_comparison_validation_roc_curve.png",
            "Figure 1. Validation ROC comparison",
        )
        self.add_picture_if_exists(
            document,
            figure_dir / "model_comparison_test_roc_curve.png",
            "Figure 2. Test ROC comparison",
        )

        document.add_heading("6. Calibration Evaluation", level=1)
        document.add_paragraph(
            "Calibration evaluates whether predicted probabilities match observed event rates. "
            "Brier score measures the mean squared difference between predicted probabilities "
            "and true binary outcomes. Lower Brier score indicates better probability accuracy."
        )
        self.add_dataframe_table(
            document,
            inputs["calibration_metrics"],
            columns=[
                "model_display_name",
                "split",
                "n_samples",
                "positive_rate",
                "brier_score",
                "n_bins_returned",
                "strategy",
            ],
            max_rows=30,
        )
        self.add_picture_if_exists(
            document,
            figure_dir / "model_comparison_validation_calibration_curve.png",
            "Figure 3. Validation calibration comparison",
        )
        self.add_picture_if_exists(
            document,
            figure_dir / "model_comparison_test_calibration_curve.png",
            "Figure 4. Test calibration comparison",
        )

        document.add_heading("7. Decision Curve Analysis", level=1)
        document.add_paragraph(
            "Decision Curve Analysis evaluates clinical net benefit across threshold probabilities. "
            "A model is considered clinically useful at a threshold if it provides higher net benefit "
            "than Treat All and Treat None strategies."
        )
        self.add_dataframe_table(
            document,
            inputs["dca_summary"],
            columns=[
                "split",
                "selected_threshold",
                "rank_at_threshold",
                "model_display_name",
                "strategy",
                "net_benefit",
                "standardized_net_benefit",
                "tp",
                "fp",
                "tn",
                "fn",
            ],
            max_rows=80,
        )
        self.add_picture_if_exists(
            document,
            figure_dir / "dca_validation_curve.png",
            "Figure 5. Validation decision curve analysis",
        )
        self.add_picture_if_exists(
            document,
            figure_dir / "dca_test_curve.png",
            "Figure 6. Test decision curve analysis",
        )

        document.add_heading("8. Interpretation Notes", level=1)
        self.add_bullet_list(
            document,
            [
                "AUC evaluates discrimination but does not evaluate calibration.",
                "Calibration is important when clinicians interpret predicted probabilities as absolute risk.",
                "DCA provides a clinical utility perspective based on threshold probabilities.",
                "Validation performance should guide model selection.",
                "Test performance should be treated as held-out final evaluation.",
                "Current results are based on simulated data and should not be interpreted as clinical evidence.",
            ],
        )

        document.add_heading("9. File Index", level=1)
        self.add_dataframe_table(
            document,
            inputs["file_index"],
            columns=["file_type", "description", "exists", "relative_path"],
            max_rows=100,
        )

        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)

        return output_path

    def build_clinical_results_docx(
        self,
        project_root: Path,
        table_dir: Path,
        figure_dir: Path,
        output_path: Path,
    ) -> Path:
        """
        Build manuscript-style clinical results report.
        """
        inputs = self.load_inputs(table_dir)

        document = Document()
        self.set_default_styles(document)

        self.add_title_page(
            document=document,
            title="Manuscript-Style Clinical Model Results Report",
            subtitle=f"Outcome: {self.config.target_name}",
            author=self.config.author,
        )

        document.add_heading("Results", level=1)

        document.add_heading("Model Development Overview", level=2)
        document.add_paragraph(
            "Three baseline machine learning models were evaluated for prediction of "
            f"{self.config.target_name}: Logistic Regression, Random Forest, and Gradient Boosting. "
            "Model selection was primarily based on validation-set performance, while the test set "
            "was preserved for held-out evaluation."
        )

        document.add_heading("Model Discrimination", level=2)
        document.add_paragraph(
            "Discrimination was assessed using AUC, accuracy, sensitivity, specificity, precision, "
            "recall, and F1 score. The validation leaderboard is shown in Table 1, and the test "
            "leaderboard is shown in Table 2."
        )

        self.add_dataframe_table(
            document,
            inputs["leaderboard_validation"],
            columns=[
                "rank_within_split",
                "model_display_name",
                "auc",
                "accuracy",
                "sensitivity",
                "specificity",
                "precision",
                "recall",
                "f1",
            ],
            max_rows=20,
            title="Table 1. Validation-set model performance",
        )

        self.add_dataframe_table(
            document,
            inputs["leaderboard_test"],
            columns=[
                "rank_within_split",
                "model_display_name",
                "auc",
                "accuracy",
                "sensitivity",
                "specificity",
                "precision",
                "recall",
                "f1",
            ],
            max_rows=20,
            title="Table 2. Test-set model performance",
        )

        self.add_picture_if_exists(
            document,
            figure_dir / "model_comparison_validation_roc_curve.png",
            "Figure 1. ROC curves for model comparison in the validation set",
        )

        self.add_picture_if_exists(
            document,
            figure_dir / "model_comparison_test_roc_curve.png",
            "Figure 2. ROC curves for model comparison in the test set",
        )

        document.add_heading("Model Calibration", level=2)
        document.add_paragraph(
            "Calibration was evaluated using calibration curves and Brier score. Lower Brier score "
            "indicates better agreement between predicted probabilities and observed outcomes."
        )

        self.add_dataframe_table(
            document,
            inputs["calibration_metrics"],
            columns=[
                "model_display_name",
                "split",
                "n_samples",
                "positive_rate",
                "brier_score",
                "n_bins_returned",
            ],
            max_rows=30,
            title="Table 3. Calibration metrics",
        )

        self.add_picture_if_exists(
            document,
            figure_dir / "model_comparison_validation_calibration_curve.png",
            "Figure 3. Calibration curves in the validation set",
        )

        self.add_picture_if_exists(
            document,
            figure_dir / "model_comparison_test_calibration_curve.png",
            "Figure 4. Calibration curves in the test set",
        )

        document.add_heading("Clinical Utility", level=2)
        document.add_paragraph(
            "Clinical utility was evaluated using Decision Curve Analysis. Net benefit was calculated "
            "across clinically relevant threshold probabilities and compared against Treat All and "
            "Treat None strategies."
        )

        self.add_dataframe_table(
            document,
            inputs["dca_summary"],
            columns=[
                "split",
                "selected_threshold",
                "rank_at_threshold",
                "model_display_name",
                "strategy",
                "net_benefit",
                "standardized_net_benefit",
            ],
            max_rows=80,
            title="Table 4. Selected-threshold DCA summary",
        )

        self.add_picture_if_exists(
            document,
            figure_dir / "dca_validation_curve.png",
            "Figure 5. Decision curve analysis in the validation set",
        )

        self.add_picture_if_exists(
            document,
            figure_dir / "dca_test_curve.png",
            "Figure 6. Decision curve analysis in the test set",
        )

        document.add_heading("Best Model Selection", level=2)
        document.add_paragraph(
            "The current best model was selected according to validation-set AUC, with sensitivity, "
            "specificity, and F1 score considered as secondary indicators."
        )

        self.add_dataframe_table(
            document,
            inputs["best_model"],
            columns=[
                "model_display_name",
                "model_name",
                "split",
                "auc",
                "accuracy",
                "sensitivity",
                "specificity",
                "precision",
                "recall",
                "f1",
                "balanced_summary_score",
                "selection_rule",
            ],
            max_rows=5,
            title="Table 5. Current best model summary",
        )

        document.add_heading("Limitations", level=2)
        self.add_bullet_list(
            document,
            [
                "The current dataset is simulated and does not represent real clinical evidence.",
                "The current models are baseline models and require external validation before clinical use.",
                "Feature importance from tree-based models should not be interpreted as causal evidence.",
                "Future steps should include SHAP interpretation, external validation, and manuscript-style statistical reporting.",
            ],
        )

        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)

        return output_path