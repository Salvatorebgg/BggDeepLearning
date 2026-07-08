"""
BggDeepLearning SHAP Word report integration module

File:
python/bggdeep/reporting/shap_word_report.py

Purpose:
1. Read existing SCI-style Word report
2. Append SHAP model interpretability section
3. Insert SHAP importance tables
4. Insert SHAP bar and beeswarm figures
5. Generate manuscript-style SHAP interpretation text
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


@dataclass
class SHAPWordReportConfig:
    """
    Configuration for SHAP Word report integration.
    """

    target_name: str = "poor_outcome"
    top_n_table: int = 15
    top_n_text: int = 5


class SHAPWordReportIntegrator:
    """
    Integrate SHAP results into existing SCI-style Word report.
    """

    def __init__(self, config: SHAPWordReportConfig) -> None:
        self.config = config
        self.table_counter_start = 5
        self.figure_counter_start = 6
        self.table_counter = self.table_counter_start
        self.figure_counter = self.figure_counter_start

    @staticmethod
    def read_csv_if_exists(path: Path) -> pd.DataFrame:
        if path.exists():
            return pd.read_csv(path)
        return pd.DataFrame()

    def load_shap_inputs(self, table_dir: Path) -> Dict[str, pd.DataFrame]:
        """
        Load SHAP importance tables.
        """
        return {
            "combined": self.read_csv_if_exists(
                table_dir / "shap_model_global_importance_all.csv"
            ),
            "random_forest": self.read_csv_if_exists(
                table_dir / "shap_random_forest_global_importance.csv"
            ),
            "gradient_boosting": self.read_csv_if_exists(
                table_dir / "shap_gradient_boosting_global_importance.csv"
            ),
        }

    @staticmethod
    def clean_feature_name(feature_name: str) -> str:
        """
        Make processed feature names easier to read in reports.
        """
        text = str(feature_name)

        replacements = {
            "num__": "",
            "cat__": "",
            "remainder__": "",
            "_": " ",
        }

        for old, new in replacements.items():
            text = text.replace(old, new)

        return text

    @staticmethod
    def format_value(value) -> str:
        if pd.isna(value):
            return ""

        if isinstance(value, float):
            if abs(value) < 0.0001 and value != 0:
                return f"{value:.2e}"
            return f"{value:.4f}"

        return str(value)

    @staticmethod
    def set_cell_border(cell, **kwargs) -> None:
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()

        tc_borders = tc_pr.first_child_found_in("w:tcBorders")
        if tc_borders is None:
            tc_borders = OxmlElement("w:tcBorders")
            tc_pr.append(tc_borders)

        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            if edge not in kwargs:
                continue

            edge_data = kwargs.get(edge)
            tag = f"w:{edge}"
            element = tc_borders.find(qn(tag))

            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)

            if edge_data is None:
                element.set(qn("w:val"), "nil")
            else:
                for key, value in edge_data.items():
                    element.set(qn(f"w:{key}"), str(value))

    @classmethod
    def clear_cell_borders(cls, cell) -> None:
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            cls.set_cell_border(cell, **{edge: None})

    @classmethod
    def apply_three_line_table_style(cls, table) -> None:
        """
        Apply three-line table style.
        """
        if not table.rows:
            return

        for row in table.rows:
            for cell in row.cells:
                cls.clear_cell_borders(cell)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        top_border = {"val": "single", "sz": "12", "space": "0", "color": "000000"}
        middle_border = {"val": "single", "sz": "8", "space": "0", "color": "000000"}
        bottom_border = {"val": "single", "sz": "12", "space": "0", "color": "000000"}

        for cell in table.rows[0].cells:
            cls.set_cell_border(cell, top=top_border, bottom=middle_border)

        for cell in table.rows[-1].cells:
            cls.set_cell_border(cell, bottom=bottom_border)

    @staticmethod
    def set_table_font(table, font_size: float = 8.2) -> None:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                        run.font.size = Pt(font_size)

    def add_table_caption(self, document: Document, title: str) -> None:
        self.table_counter += 1

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(f"Table {self.table_counter}. {title}")
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)

    def add_figure_caption(self, document: Document, title: str) -> None:
        self.figure_counter += 1

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"Figure {self.figure_counter}. {title}")
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)

    def add_paragraph(self, document: Document, text: str) -> None:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10.5)

    def prepare_shap_table(self, df: pd.DataFrame) -> pd.DataFrame:
        """
        Prepare SHAP table for Word output.
        """
        if df.empty:
            return df

        display_df = df.copy()

        if "feature_name" in display_df.columns:
            display_df["feature_label"] = display_df["feature_name"].apply(
                self.clean_feature_name
            )

        columns = [
            "rank",
            "feature_label",
            "mean_abs_shap",
        ]

        existing_columns = [
            column for column in columns
            if column in display_df.columns
        ]

        display_df = display_df[existing_columns].head(self.config.top_n_table)

        rename_map = {
            "rank": "Rank",
            "feature_label": "Feature",
            "mean_abs_shap": "Mean absolute SHAP",
        }

        display_df = display_df.rename(columns=rename_map)

        return display_df

    def add_sci_table(
        self,
        document: Document,
        df: pd.DataFrame,
        title: str,
        note: str | None = None,
    ) -> None:
        """
        Add SCI-style SHAP table.
        """
        self.add_table_caption(document, title)

        if df.empty:
            document.add_paragraph("No available data.")
            return

        table = document.add_table(rows=1, cols=len(df.columns))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        header_cells = table.rows[0].cells
        for i, column in enumerate(df.columns):
            header_cells[i].text = str(column)

        for _, row in df.iterrows():
            cells = table.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = self.format_value(value)

        self.apply_three_line_table_style(table)
        self.set_table_font(table, font_size=8.2)

        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True

        if note:
            p = document.add_paragraph()
            r = p.add_run(f"Note. {note}")
            r.font.name = "Times New Roman"
            r.font.size = Pt(8)

        document.add_paragraph("")

    def add_figure_if_exists(
        self,
        document: Document,
        image_path: Path,
        title: str,
        width_inches: float = 6.2,
    ) -> None:
        """
        Add SHAP figure if it exists.
        """
        if not image_path.exists():
            p = document.add_paragraph()
            r = p.add_run(f"Figure file not found: {image_path}")
            r.italic = True
            r.font.size = Pt(9)
            return

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(width_inches))

        self.add_figure_caption(document, title)
        document.add_paragraph("")

    def build_interpretation_sentence(
        self,
        model_display_name: str,
        importance_df: pd.DataFrame,
    ) -> str:
        """
        Build one manuscript-style interpretation sentence for a model.
        """
        if importance_df.empty:
            return f"For {model_display_name}, SHAP results were not available."

        top_df = importance_df.head(self.config.top_n_text).copy()

        top_features = [
            self.clean_feature_name(value)
            for value in top_df["feature_name"].tolist()
        ]

        if len(top_features) == 1:
            feature_text = top_features[0]
        else:
            feature_text = ", ".join(top_features[:-1]) + f", and {top_features[-1]}"

        return (
            f"For {model_display_name}, global SHAP analysis identified {feature_text} "
            f"as the leading contributors to model predictions for {self.config.target_name}. "
            "These variables had the largest mean absolute SHAP values, indicating stronger "
            "overall influence on the model output."
        )

    def build_shap_interpretation_text(
        self,
        shap_inputs: Dict[str, pd.DataFrame],
    ) -> str:
        """
        Build manuscript-style SHAP interpretation text.
        """
        rf_text = self.build_interpretation_sentence(
            "Random Forest",
            shap_inputs["random_forest"],
        )

        gb_text = self.build_interpretation_sentence(
            "Gradient Boosting",
            shap_inputs["gradient_boosting"],
        )

        lines = []
        lines.append("Model interpretability")
        lines.append("")
        lines.append(
            "Global model interpretability was assessed using SHAP analysis for the tree-based models. "
            "Mean absolute SHAP values were used to rank features according to their overall influence "
            "on model predictions."
        )
        lines.append("")
        lines.append(rf_text)
        lines.append("")
        lines.append(gb_text)
        lines.append("")
        lines.append(
            "SHAP values explain model behavior rather than causal relationships. Therefore, the identified "
            "features should be interpreted as variables that the trained models relied on for prediction, "
            "not as confirmed causal determinants of the clinical outcome."
        )

        return "\n".join(lines)

    def append_shap_section(
        self,
        base_docx: Path,
        output_docx: Path,
        table_dir: Path,
        figure_dir: Path,
        report_dir: Path,
    ) -> Dict[str, Path]:
        """
        Append SHAP section to existing SCI Word report.
        """
        if not base_docx.exists():
            raise FileNotFoundError(
                f"Base SCI Word report was not found: {base_docx}\n"
                "Please run: python scripts\\python\\export_sci_word_report.py"
            )

        shap_inputs = self.load_shap_inputs(table_dir)

        if shap_inputs["combined"].empty:
            raise FileNotFoundError(
                "SHAP combined importance table was not found or empty.\n"
                "Please run: python scripts\\python\\run_shap_explainability.py"
            )

        document = Document(str(base_docx))

        document.add_page_break()

        document.add_heading("Model interpretability", level=2)

        interpretation_text = self.build_shap_interpretation_text(shap_inputs)

        for paragraph_text in interpretation_text.split("\n"):
            if paragraph_text.strip() == "":
                continue

            if paragraph_text.strip() == "Model interpretability":
                continue

            self.add_paragraph(document, paragraph_text)

        rf_table = self.prepare_shap_table(shap_inputs["random_forest"])
        gb_table = self.prepare_shap_table(shap_inputs["gradient_boosting"])

        self.add_sci_table(
            document=document,
            df=rf_table,
            title="Top SHAP features for the Random Forest model",
            note="Features are ranked by mean absolute SHAP value on the processed test set.",
        )

        self.add_sci_table(
            document=document,
            df=gb_table,
            title="Top SHAP features for the Gradient Boosting model",
            note="Features are ranked by mean absolute SHAP value on the processed test set.",
        )

        self.add_figure_if_exists(
            document=document,
            image_path=figure_dir / "shap_random_forest_bar.png",
            title="Global SHAP feature importance for the Random Forest model.",
        )

        self.add_figure_if_exists(
            document=document,
            image_path=figure_dir / "shap_gradient_boosting_bar.png",
            title="Global SHAP feature importance for the Gradient Boosting model.",
        )

        self.add_figure_if_exists(
            document=document,
            image_path=figure_dir / "shap_random_forest_beeswarm.png",
            title="SHAP beeswarm plot for the Random Forest model.",
        )

        self.add_figure_if_exists(
            document=document,
            image_path=figure_dir / "shap_gradient_boosting_beeswarm.png",
            title="SHAP beeswarm plot for the Gradient Boosting model.",
        )

        output_docx.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_docx)

        interpretation_file = report_dir / "shap_interpretation_text.txt"
        interpretation_file.write_text(interpretation_text, encoding="utf-8")

        return {
            "shap_docx": output_docx,
            "interpretation_text": interpretation_file,
        }