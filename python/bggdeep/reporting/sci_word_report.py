"""
BggDeepLearning SCI-style Word report module

File:
python/bggdeep/reporting/sci_word_report.py

Purpose:
1. Generate a SCI-style clinical results DOCX report
2. Use three-line tables
3. Add automatic table and figure numbering
4. Insert ROC, calibration, and DCA figures
5. Produce a cleaner manuscript-like Results section
"""

from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Dict, List

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


@dataclass
class SCIWordReportConfig:
    project_name: str = "BggDeepLearning"
    target_name: str = "poor_outcome"
    report_title: str = "Clinical Machine Learning Model Results"
    author: str = "BggDeepLearning Team"


class SCIWordReportBuilder:
    """
    Build SCI-style Word report.
    """

    def __init__(self, config: SCIWordReportConfig) -> None:
        self.config = config
        self.table_counter = 0
        self.figure_counter = 0

    @staticmethod
    def read_csv_if_exists(path: Path) -> pd.DataFrame:
        if path.exists():
            return pd.read_csv(path)
        return pd.DataFrame()

    def load_inputs(self, table_dir: Path) -> Dict[str, pd.DataFrame]:
        return {
            "best_model": self.read_csv_if_exists(table_dir / "model_best_model_summary.csv"),
            "leaderboard_validation": self.read_csv_if_exists(table_dir / "model_leaderboard_validation.csv"),
            "leaderboard_test": self.read_csv_if_exists(table_dir / "model_leaderboard_test.csv"),
            "roc_summary": self.read_csv_if_exists(table_dir / "model_comparison_roc_summary.csv"),
            "calibration_metrics": self.read_csv_if_exists(table_dir / "model_calibration_metrics.csv"),
            "dca_summary": self.read_csv_if_exists(table_dir / "dca_selected_threshold_summary.csv"),
            "file_index": self.read_csv_if_exists(table_dir / "comprehensive_model_report_file_index.csv"),
        }

    @staticmethod
    def set_page_layout(document: Document) -> None:
        section = document.sections[0]
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    @staticmethod
    def set_default_styles(document: Document) -> None:
        styles = document.styles

        normal = styles["Normal"]
        normal.font.name = "Times New Roman"
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        normal.font.size = Pt(10.5)

        for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
            if style_name in styles:
                styles[style_name].font.name = "Times New Roman"
                styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

        if "Heading 1" in styles:
            styles["Heading 1"].font.size = Pt(14)
            styles["Heading 1"].font.bold = True

        if "Heading 2" in styles:
            styles["Heading 2"].font.size = Pt(12)
            styles["Heading 2"].font.bold = True

    @staticmethod
    def add_title(document: Document, title: str, subtitle: str, author: str) -> None:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(18)

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(subtitle)
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(author)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"Generated time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        r.font.name = "Times New Roman"
        r.font.size = Pt(9)

        document.add_paragraph("")

    @staticmethod
    def set_cell_border(cell, **kwargs) -> None:
        """
        Set cell borders.

        Usage:
        set_cell_border(cell, top={"val": "single", "sz": "12", "color": "000000"})
        """
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()

        tc_borders = tc_pr.first_child_found_in("w:tcBorders")
        if tc_borders is None:
            tc_borders = OxmlElement("w:tcBorders")
            tc_pr.append(tc_borders)

        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            if edge not in kwargs:
                continue

            edge_data = kwargs.get(edge)
            tag = f"w:{edge}"
            element = tc_borders.find(qn(tag))

            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)

            if edge_data is None:
                element.set(qn("w:val"), "nil")
            else:
                for key, value in edge_data.items():
                    element.set(qn(f"w:{key}"), str(value))

    @classmethod
    def clear_cell_borders(cls, cell) -> None:
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            cls.set_cell_border(cell, **{edge: None})

    @classmethod
    def apply_three_line_table_style(cls, table) -> None:
        """
        Apply SCI-style three-line table:
        1. top border of header row
        2. bottom border of header row
        3. bottom border of last row
        """
        if not table.rows:
            return

        for row in table.rows:
            for cell in row.cells:
                cls.clear_cell_borders(cell)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        top_border = {"val": "single", "sz": "12", "space": "0", "color": "000000"}
        middle_border = {"val": "single", "sz": "8", "space": "0", "color": "000000"}
        bottom_border = {"val": "single", "sz": "12", "space": "0", "color": "000000"}

        for cell in table.rows[0].cells:
            cls.set_cell_border(cell, top=top_border, bottom=middle_border)

        for cell in table.rows[-1].cells:
            cls.set_cell_border(cell, bottom=bottom_border)

    @staticmethod
    def mark_first_row_as_header(table) -> None:
        """
        Repeat header row on new pages.
        """
        if not table.rows:
            return

        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)

    @staticmethod
    def set_table_font(table, font_size: float = 8.5) -> None:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                        run.font.size = Pt(font_size)

    @staticmethod
    def format_value(value) -> str:
        if pd.isna(value):
            return ""

        if isinstance(value, float):
            if abs(value) < 0.0001 and value != 0:
                return f"{value:.2e}"
            return f"{value:.4f}"

        return str(value)

    def add_table_caption(self, document: Document, title: str) -> None:
        self.table_counter += 1
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(f"Table {self.table_counter}. {title}")
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)

    def add_figure_caption(self, document: Document, title: str) -> None:
        self.figure_counter += 1
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"Figure {self.figure_counter}. {title}")
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)

    def add_sci_table(
        self,
        document: Document,
        df: pd.DataFrame,
        title: str,
        columns: List[str] | None = None,
        max_rows: int = 20,
        note: str | None = None,
    ) -> None:
        self.add_table_caption(document, title)

        if df.empty:
            document.add_paragraph("No available data.")
            return

        display_df = df.copy()

        if columns is not None:
            existing_columns = [col for col in columns if col in display_df.columns]
            if existing_columns:
                display_df = display_df[existing_columns]

        display_df = display_df.head(max_rows)

        table = document.add_table(rows=1, cols=len(display_df.columns))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        header_cells = table.rows[0].cells
        for i, column in enumerate(display_df.columns):
            header_cells[i].text = str(column)

        for _, row in display_df.iterrows():
            cells = table.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = self.format_value(value)

        self.apply_three_line_table_style(table)
        self.mark_first_row_as_header(table)
        self.set_table_font(table, font_size=8.2)

        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True

        if note:
            p = document.add_paragraph()
            r = p.add_run(f"Note. {note}")
            r.font.name = "Times New Roman"
            r.font.size = Pt(8)

        document.add_paragraph("")

    def add_sci_figure(
        self,
        document: Document,
        image_path: Path,
        title: str,
        width_inches: float = 6.2,
    ) -> None:
        if not image_path.exists():
            p = document.add_paragraph()
            r = p.add_run(f"Figure file not found: {image_path}")
            r.italic = True
            r.font.size = Pt(9)
            return

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(width_inches))

        self.add_figure_caption(document, title)
        document.add_paragraph("")

    @staticmethod
    def add_paragraph(document: Document, text: str) -> None:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10.5)

    def build_sci_results_docx(
        self,
        project_root: Path,
        table_dir: Path,
        figure_dir: Path,
        output_path: Path,
    ) -> Path:
        inputs = self.load_inputs(table_dir)

        validation_df = inputs["leaderboard_validation"]
        test_df = inputs["leaderboard_test"]
        calibration_df = inputs["calibration_metrics"]
        dca_df = inputs["dca_summary"]
        best_model_df = inputs["best_model"]

        if not calibration_df.empty and "brier_score" in calibration_df.columns:
            calibration_df = calibration_df.sort_values(
                by=["split", "brier_score"],
                ascending=[True, True],
            )

        if not dca_df.empty:
            dca_df = dca_df[
                dca_df["selected_threshold"].isin([0.10, 0.20, 0.30, 0.50])
            ].copy()

        document = Document()
        self.set_default_styles(document)
        self.set_page_layout(document)

        self.add_title(
            document=document,
            title=self.config.report_title,
            subtitle=f"Outcome: {self.config.target_name}",
            author=self.config.author,
        )

        document.add_heading("Results", level=1)

        document.add_heading("Model development and evaluation overview", level=2)
        self.add_paragraph(
            document,
            "Three baseline machine learning models were evaluated for prediction of "
            f"{self.config.target_name}: Logistic Regression, Random Forest, and Gradient Boosting. "
            "Model selection was based on validation-set performance, while the test set was "
            "reserved for held-out evaluation. The present report was generated from simulated "
            "clinical data and is intended for engineering demonstration."
        )

        document.add_heading("Discrimination performance", level=2)
        self.add_paragraph(
            document,
            "Model discrimination was evaluated using AUC, accuracy, sensitivity, specificity, "
            "precision, recall, and F1 score. The validation-set and test-set rankings are shown "
            "in the following tables."
        )

        performance_columns = [
            "rank_within_split",
            "model_display_name",
            "auc",
            "accuracy",
            "sensitivity",
            "specificity",
            "precision",
            "recall",
            "f1",
        ]

        self.add_sci_table(
            document=document,
            df=validation_df,
            title="Validation-set performance of candidate models",
            columns=performance_columns,
            max_rows=20,
            note="Models are ranked using validation-set AUC. Sensitivity, specificity, and F1 score are included for clinical interpretation.",
        )

        self.add_sci_table(
            document=document,
            df=test_df,
            title="Test-set performance of candidate models",
            columns=performance_columns,
            max_rows=20,
            note="The test set is used as held-out final evaluation and should not be repeatedly used for model selection.",
        )

        self.add_sci_figure(
            document=document,
            image_path=figure_dir / "model_comparison_validation_roc_curve.png",
            title="ROC curves for candidate models in the validation set.",
        )

        self.add_sci_figure(
            document=document,
            image_path=figure_dir / "model_comparison_test_roc_curve.png",
            title="ROC curves for candidate models in the test set.",
        )

        document.add_heading("Calibration performance", level=2)
        self.add_paragraph(
            document,
            "Calibration was evaluated using calibration curves and Brier score. Brier score "
            "measures the mean squared difference between predicted probabilities and true "
            "binary outcomes; lower values indicate better probability accuracy."
        )

        self.add_sci_table(
            document=document,
            df=calibration_df,
            title="Calibration metrics of candidate models",
            columns=[
                "model_display_name",
                "split",
                "n_samples",
                "positive_rate",
                "brier_score",
                "n_bins_returned",
                "strategy",
            ],
            max_rows=30,
            note="Lower Brier score indicates better overall calibration and probability accuracy.",
        )

        self.add_sci_figure(
            document=document,
            image_path=figure_dir / "model_comparison_validation_calibration_curve.png",
            title="Calibration curves for candidate models in the validation set.",
        )

        self.add_sci_figure(
            document=document,
            image_path=figure_dir / "model_comparison_test_calibration_curve.png",
            title="Calibration curves for candidate models in the test set.",
        )

        document.add_heading("Clinical utility", level=2)
        self.add_paragraph(
            document,
            "Decision Curve Analysis was used to evaluate clinical net benefit across threshold "
            "probabilities. A model is potentially useful at a given threshold when it provides "
            "higher net benefit than Treat All and Treat None strategies."
        )

        self.add_sci_table(
            document=document,
            df=dca_df,
            title="Selected-threshold Decision Curve Analysis summary",
            columns=[
                "split",
                "selected_threshold",
                "rank_at_threshold",
                "model_display_name",
                "strategy",
                "net_benefit",
                "standardized_net_benefit",
            ],
            max_rows=80,
            note="Only selected clinically interpretable threshold probabilities are displayed.",
        )

        self.add_sci_figure(
            document=document,
            image_path=figure_dir / "dca_validation_curve.png",
            title="Decision Curve Analysis in the validation set.",
        )

        self.add_sci_figure(
            document=document,
            image_path=figure_dir / "dca_test_curve.png",
            title="Decision Curve Analysis in the test set.",
        )

        document.add_heading("Best model selection", level=2)
        self.add_paragraph(
            document,
            "The current best model was selected according to validation-set AUC, with sensitivity, "
            "specificity, and F1 score considered as secondary indicators."
        )

        self.add_sci_table(
            document=document,
            df=best_model_df,
            title="Current best model summary",
            columns=[
                "model_display_name",
                "model_name",
                "split",
                "auc",
                "accuracy",
                "sensitivity",
                "specificity",
                "precision",
                "recall",
                "f1",
                "balanced_summary_score",
            ],
            max_rows=5,
            note="The selected model should still require external validation before clinical deployment.",
        )

        document.add_heading("Limitations", level=2)
        limitations = [
            "The current dataset is simulated and should not be interpreted as real-world clinical evidence.",
            "The current models are baseline models and require further validation before clinical use.",
            "Tree-based feature importance should not be interpreted as causal evidence.",
            "Future work should include SHAP interpretation, external validation, and manuscript-level statistical reporting.",
        ]

        for item in limitations:
            document.add_paragraph(item, style="List Bullet")

        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)

        return output_path